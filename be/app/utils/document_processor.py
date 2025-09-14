import os
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging
import base64
from io import BytesIO
from ..core.config import Config

# LangChain text splitter
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter, CharacterTextSplitter
except ImportError:
    RecursiveCharacterTextSplitter = None
    MarkdownHeaderTextSplitter = None
    CharacterTextSplitter = None

# PDF processing
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

# DOCX processing
try:
    from docx import Document
except ImportError:
    Document = None

# Image processing
try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, data_dir: str = "data/thutuccongdan", openai_service=None):
        """
        Initialize document processor
        
        Args:
            data_dir: Directory containing files
            openai_service: OpenAI service instance for LLM-based extraction
        """
        self.data_dir = data_dir
        self.chunk_size = Config.CHUNK_SIZE
        self.chunk_overlap = Config.CHUNK_OVERLAP
        self.separators = Config.CHUNK_SEPARATORS
        self.openai_service = openai_service
        
        # Initialize LangChain fixed-size text splitter (CharacterTextSplitter for consistency)
        if CharacterTextSplitter:
            self.text_splitter = CharacterTextSplitter(
                chunk_size=Config.CHUNK_SIZE,
                chunk_overlap=Config.CHUNK_OVERLAP,
                separator="\n\n",  # Split by paragraphs primarily
                length_function=len
            )
            logger.info(f"🔧 [PROCESSOR] Initialized LangChain CharacterTextSplitter (fixed-size) with chunk_size={Config.CHUNK_SIZE}, overlap={Config.CHUNK_OVERLAP}")
        elif RecursiveCharacterTextSplitter:
            # Fallback to RecursiveCharacterTextSplitter with limited separators
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=Config.CHUNK_SIZE,
                chunk_overlap=Config.CHUNK_OVERLAP,
                separators=["\n\n", "\n", ". ", " "],  # Limited separators for more consistent sizes
                length_function=len,
                is_separator_regex=False
            )
            logger.info(f"🔧 [PROCESSOR] Initialized RecursiveCharacterTextSplitter as fallback with chunk_size={Config.CHUNK_SIZE}, overlap={Config.CHUNK_OVERLAP}")
        else:
            self.text_splitter = None
            logger.warning("⚠️ [PROCESSOR] LangChain not available, falling back to simple chunking")
            
        # Initialize Markdown header splitter for header preservation
        if MarkdownHeaderTextSplitter:
            self.md_header_splitter = MarkdownHeaderTextSplitter(
                headers_to_split_on=[
                    ("#", "Header 1"),
                    ("##", "Header 2"), 
                    ("###", "Header 3"),
                    ("####", "Header 4"),
                    ("#####", "Header 5"),
                    ("######", "Header 6"),
                ]
            )
            logger.info("🔧 [PROCESSOR] Initialized Markdown header splitter for context preservation")
        else:
            self.md_header_splitter = None
        
        # Supported file types and their processors
        self.processors = {
            '.pdf': self._extract_pdf_content,
            '.docx': self._extract_docx_content,
            '.doc': self._extract_docx_content,
            '.txt': self._extract_text_content,
            '.md': self._extract_markdown_content,
            '.png': self._extract_image_content,
            '.jpg': self._extract_image_content,
            '.jpeg': self._extract_image_content
        }
        
    def process_uploaded_file(self, file_path: str, filename: str) -> Optional[Dict[str, Any]]:
        """
        Process an uploaded file and extract its content
        
        Args:
            file_path: Path to the uploaded file
            filename: Original filename
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        logger.info(f"🔄 [PROCESSOR] Starting to process file: {filename}")
        logger.info(f"📍 [PROCESSOR] File path: {file_path}")
        logger.info(f"📁 [PROCESSOR] File exists: {os.path.exists(file_path)}")
        
        try:
            file_extension = os.path.splitext(filename)[1].lower()
            logger.info(f"📎 [PROCESSOR] File extension: {file_extension}")
            
            if file_extension not in self.processors:
                logger.warning(f"❌ [PROCESSOR] Unsupported file type: {file_extension}")
                logger.info(f"✅ [PROCESSOR] Supported types: {list(self.processors.keys())}")
                return None
            
            logger.info(f"🔧 [PROCESSOR] Using processor for {file_extension}")
            # Extract content using appropriate processor
            processor = self.processors[file_extension]
            content = processor(file_path)
            
            logger.info(f"📝 [PROCESSOR] Content extracted, length: {len(content) if content else 0} characters")
            
            if not content:
                logger.warning(f"⚠️ [PROCESSOR] No content extracted from file: {filename}")
                return None
            
            # Create document object
            doc = {
                "file_name": filename,
                "file_path": file_path,
                "file_type": file_extension,
                "title": self._extract_title_from_filename(filename),
                "content": content,
                "processed_at": None  # Will be set by caller
            }
            
            logger.info(f"✅ [PROCESSOR] Successfully processed file: {filename}")
            logger.info(f"📊 [PROCESSOR] Document stats - Title: '{doc['title']}', Content: {len(content)} chars")
            return doc
            
        except Exception as e:
            logger.error(f"❌ [PROCESSOR] Failed to process file {filename}: {e}", exc_info=True)
            return None
    
    def _extract_title_from_filename(self, filename: str) -> str:
        """Extract title from filename"""
        # Remove extension and replace underscores/hyphens with spaces
        title = os.path.splitext(filename)[0]
        title = re.sub(r'[_-]', ' ', title)
        return title.strip()
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """Extract text content from PDF file"""
        if not PyPDF2:
            logger.error("PyPDF2 not available. Install PyPDF2 to process PDF files.")
            return ""
        
        try:
            content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content += f"\n\n--- Page {page_num + 1} ---\n"
                            content += page_text
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"Error reading PDF file: {e}")
            return ""
    
    def _extract_docx_content(self, file_path: str) -> str:
        """Extract text content from DOCX file"""
        if not Document:
            logger.error("python-docx not available. Install python-docx to process DOCX files.")
            return ""
        
        try:
            doc = Document(file_path)
            content = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content += " | ".join(row_text) + "\n"
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"Error reading DOCX file: {e}")
            return ""
    
    def _extract_text_content(self, file_path: str) -> str:
        """Extract content from text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors='replace'
            with open(file_path, 'rb') as file:
                raw_content = file.read()
                return raw_content.decode('utf-8', errors='replace')
                
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return ""
    
    def _extract_markdown_content(self, file_path: str) -> str:
        """Extract content from markdown file"""
        return self._extract_text_content(file_path)
    
    def _extract_image_content(self, file_path: str) -> str:
        """Extract text content from image using OCR and LLM vision"""
        content = ""
        
        # First try OCR if available
        if Image and pytesseract:
            try:
                with Image.open(file_path) as img:
                    ocr_text = pytesseract.image_to_string(img, lang='vie+eng')
                    if ocr_text.strip():
                        content += "=== OCR Extracted Text ===\n" + ocr_text.strip() + "\n\n"
            except Exception as e:
                logger.warning(f"OCR extraction failed: {e}")
        
        # Then use LLM vision if OpenAI service is available
        if self.openai_service:
            try:
                llm_content = self._extract_image_content_with_llm(file_path)
                if llm_content:
                    content += "=== AI Vision Analysis ===\n" + llm_content
            except Exception as e:
                logger.warning(f"LLM vision extraction failed: {e}")
        
        return content.strip()
    
    def _extract_image_content_with_llm(self, file_path: str) -> str:
        """Extract content from image using LLM vision capabilities"""
        try:
            # Encode image to base64
            with open(file_path, 'rb') as img_file:
                img_bytes = img_file.read()
                img_base64 = base64.b64encode(img_bytes).decode('utf-8')
            
            # Prepare messages for GPT-4V
            messages = [
                {
                    "role": "user", 
                    "content": [
                        {
                            "type": "text",
                            "text": "Hãy phân tích hình ảnh này và trích xuất tất cả văn bản có thể đọc được. Nếu có bảng biểu, hãy mô tả cấu trúc và nội dung. Nếu có biểu đồ hoặc sơ đồ, hãy giải thích ý nghĩa. Trả lời bằng tiếng Việt."
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{img_base64}"
                            }
                        }
                    ]
                }
            ]
            
            # Use GPT-4V for image analysis
            response = self.openai_service.chat_completion(
                messages, 
                model="gpt-4o", 
                max_tokens=1500
            )
            
            return response.strip()
            
        except Exception as e:
            logger.error(f"Error using LLM for image analysis: {e}")
            return ""
        
    def read_markdown_files(self) -> List[Dict[str, Any]]:
        """
        Read all markdown files and extract content
        
        Returns:
            List of document dictionaries
        """
        documents = []
        data_path = Path(self.data_dir)
        
        if not data_path.exists():
            logger.error(f"Data directory not found: {self.data_dir}")
            return documents
        
        # Get all markdown files
        md_files = list(data_path.glob("*.md"))
        logger.info(f"Found {len(md_files)} markdown files")
        
        for md_file in md_files:
            try:
                with open(md_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Extract document metadata
                title, sections = self._parse_markdown_content(content)
                
                # Create document object
                doc = {
                    "file_name": md_file.name,
                    "title": title,
                    "content": content,
                    "sections": sections
                }
                
                documents.append(doc)
                logger.info(f"Processed file: {md_file.name}")
                
            except Exception as e:
                logger.error(f"Failed to read file {md_file}: {e}")
        
        return documents
    
    def _parse_markdown_content(self, content: str) -> tuple:
        """
        Parse markdown content to extract title and sections
        
        Args:
            content: Raw markdown content
            
        Returns:
            Tuple of (title, sections)
        """
        lines = content.split('\n')
        title = "Untitled"
        sections = []
        current_section = ""
        current_section_title = ""
        
        for line in lines:
            # Extract main title (# header)
            if line.startswith('# ') and not title or title == "Untitled":
                title = line[2:].strip()
            
            # Extract sections (## headers)
            elif line.startswith('## '):
                # Save previous section
                if current_section.strip():
                    sections.append({
                        "title": current_section_title,
                        "content": current_section.strip()
                    })
                
                # Start new section
                current_section_title = line[3:].strip()
                current_section = ""
            
            else:
                current_section += line + '\n'
        
        # Add last section
        if current_section.strip():
            sections.append({
                "title": current_section_title,
                "content": current_section.strip()
            })
        
        return title, sections
    
    def chunk_documents(self, documents: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Split documents into chunks for embedding
        
        Args:
            documents: List of document dictionaries
            
        Returns:
            List of chunked documents
        """
        chunked_docs = []
        
        for doc in documents:
            # Process each section separately
            for section_idx, section in enumerate(doc["sections"]):
                section_content = section["content"]
                section_title = section["title"]
                
                # Split section into chunks
                chunks = self._split_text(section_content)
                
                for chunk_idx, chunk in enumerate(chunks):
                    chunked_doc = {
                        "file_name": doc["file_name"],
                        "chunk_id": section_idx * 100 + chunk_idx,  # Unique chunk ID
                        "content": chunk,
                        "title": doc["title"],
                        "section": section_title
                    }
                    chunked_docs.append(chunked_doc)
            
            # Also create chunks for the full document
            full_content_chunks = self._split_text(doc["content"], ".md")  # Assume markdown for legacy method
            for chunk_idx, chunk in enumerate(full_content_chunks):
                chunked_doc = {
                    "file_name": doc["file_name"],
                    "chunk_id": 9000 + chunk_idx,  # Different range for full document chunks
                    "content": chunk,
                    "title": doc["title"],
                    "section": "Full Document"
                }
                chunked_docs.append(chunked_doc)
        
        logger.info(f"Created {len(chunked_docs)} chunks from {len(documents)} documents")
        return chunked_docs
    
    def _split_text(self, text: str, file_type: str = None) -> List[str]:
        """
        Split text using LangChain's fixed-size chunking for consistent chunks
        
        Args:
            text: Text to split
            file_type: File extension to determine header preservation
            
        Returns:
            List of fixed-size text chunks
        """
        logger.debug(f"📝 [CHUNKER] Starting LangChain fixed-size split for {len(text)} characters")
        logger.debug(f"⚙️ [CHUNKER] Using CHUNK_SIZE={Config.CHUNK_SIZE}, OVERLAP={Config.CHUNK_OVERLAP}")
        logger.debug(f"🔧 [CHUNKER] File type: {file_type}")
        
        try:
            if not text.strip():
                return []
            
            # Use LangChain's fixed-size chunking
            return self._langchain_fixed_size_split(text, file_type or "")
                
        except Exception as e:
            logger.warning(f"🔄 LangChain fixed-size chunking failed: {e}. Using simple fallback.")
            return self._simple_split(text)
    
    def _split_markdown_with_headers(self, text: str) -> List[str]:
        """Split markdown text preserving header context"""
        logger.info("📋 [CHUNKER] Using Markdown header-aware splitting")
        
        try:
            # First split by headers to get sections with metadata
            md_header_splits = self.md_header_splitter.split_text(text)
            
            final_chunks = []
            
            for doc in md_header_splits:
                # Get header context from metadata
                header_context = self._build_header_context(doc.metadata)
                
                # Split the content if it's too large
                if len(doc.page_content) <= self.chunk_size:
                    # Small enough, keep as is with header context
                    chunk_with_context = header_context + doc.page_content if header_context else doc.page_content
                    final_chunks.append(chunk_with_context.strip())
                else:
                    # Too large, need to split further
                    sub_chunks = self.text_splitter.split_text(doc.page_content)
                    
                    for sub_chunk in sub_chunks:
                        if sub_chunk.strip():
                            # Add header context to each sub-chunk
                            chunk_with_context = header_context + sub_chunk if header_context else sub_chunk
                            final_chunks.append(chunk_with_context.strip())
            
            logger.info(f"🎯 [CHUNKER] Markdown header splitting generated {len(final_chunks)} chunks")
            return final_chunks
            
        except Exception as e:
            logger.error(f"💥 [CHUNKER] Markdown header splitting failed: {e}")
            return self._standard_split(text)
    
    def _split_text_with_header_context(self, text: str) -> List[str]:
        """Split text while preserving header context for non-markdown files"""
        logger.info("📄 [CHUNKER] Using header context preservation for structured text")
        
        try:
            # Extract headers and their positions
            headers = self._extract_headers(text)
            
            # Standard split first
            chunks = self.text_splitter.split_text(text)
            
            # Add header context to each chunk
            enhanced_chunks = []
            
            for chunk in chunks:
                if chunk.strip():
                    # Find relevant headers for this chunk
                    chunk_start = text.find(chunk[:50])  # Find approximate position
                    relevant_headers = self._find_relevant_headers(headers, chunk_start)
                    
                    # Build header context
                    header_context = self._build_text_header_context(relevant_headers)
                    
                    # Combine header context with chunk
                    if header_context:
                        enhanced_chunk = header_context + "\n\n" + chunk
                    else:
                        enhanced_chunk = chunk
                    
                    enhanced_chunks.append(enhanced_chunk.strip())
            
            logger.info(f"🎯 [CHUNKER] Header context splitting generated {len(enhanced_chunks)} chunks")
            return enhanced_chunks
            
        except Exception as e:
            logger.error(f"💥 [CHUNKER] Header context splitting failed: {e}")
            return self._standard_split(text)
    
    def _langchain_fixed_size_split(self, text: str, file_type: str) -> List[str]:
        """LangChain-based fixed-size chunking with consistent chunk sizes"""
        try:
            # For markdown files with header preservation
            if file_type == '.md' and Config.PRESERVE_HEADERS:
                return self._langchain_markdown_fixed_split(text)
            
            # Standard fixed-size splitting using LangChain
            if not self.text_splitter:
                logger.warning("🔄 LangChain splitter not available, using fallback")
                return self._simple_split(text)
            
            chunks = self.text_splitter.split_text(text)
            chunks = [chunk.strip() for chunk in chunks if chunk.strip()]
            
            logger.info(f"📊 LangChain fixed-size split: {len(chunks)} chunks created")
            for i, chunk in enumerate(chunks):
                logger.debug(f"   Chunk {i+1}: {len(chunk)} chars")
            
            return chunks
            
        except Exception as e:
            logger.error(f"❌ LangChain fixed-size splitting failed: {e}")
            return self._simple_split(text)
    
    def _langchain_markdown_fixed_split(self, text: str) -> List[str]:
        """Fixed-size chunking for markdown with header preservation"""
        try:
            # Extract document headers for context
            headers_context = self._extract_document_headers(text)
            
            # Use LangChain's CharacterTextSplitter for consistent sizes
            chunks = self.text_splitter.split_text(text)
            chunks = [chunk.strip() for chunk in chunks if chunk.strip()]
            
            # Add header context to each chunk if available
            if headers_context:
                enhanced_chunks = []
                for chunk in chunks:
                    # Add headers to maintain context
                    enhanced_chunk = headers_context + "\n\n" + chunk
                    enhanced_chunks.append(enhanced_chunk)
                chunks = enhanced_chunks
            
            logger.info(f"📊 LangChain markdown fixed-split: {len(chunks)} chunks with headers")
            for i, chunk in enumerate(chunks):
                logger.debug(f"   Chunk {i+1}: {len(chunk)} chars")
            
            return chunks
            
        except Exception as e:
            logger.error(f"❌ LangChain markdown fixed-splitting failed: {e}")
            return self.text_splitter.split_text(text) if self.text_splitter else self._simple_split(text)
    
    def _fixed_size_split(self, text: str, file_type: str) -> List[str]:
        """Fixed-size chunking with optional header preservation"""
        try:
            chunk_size = Config.CHUNK_SIZE
            chunk_overlap = Config.CHUNK_OVERLAP
            
            # Extract headers if markdown and preserve_headers is enabled
            headers_context = ""
            if file_type == '.md' and Config.PRESERVE_HEADERS:
                headers_context = self._extract_document_headers(text)
            
            chunks = []
            start = 0
            text_length = len(text)
            
            while start < text_length:
                # Calculate end position
                end = start + chunk_size
                
                # If this is the last chunk and it's very small, merge with previous
                if end >= text_length:
                    chunk_text = text[start:]
                    if len(chunk_text) < chunk_size * 0.3 and chunks:  # Less than 30% of chunk_size
                        # Merge with last chunk
                        chunks[-1] = chunks[-1] + "\n\n" + chunk_text
                    else:
                        # Add header context if needed
                        if headers_context and file_type == '.md':
                            chunk_text = headers_context + "\n\n" + chunk_text
                        chunks.append(chunk_text)
                    break
                
                # Find a good break point within the overlap zone
                break_point = end
                
                # Look for natural break points in the overlap zone
                search_start = max(start + chunk_size - chunk_overlap, start + chunk_size // 2)
                search_end = min(end + chunk_overlap, text_length)
                
                # Try to find good break points in order of preference
                break_candidates = []
                
                # Look for paragraph breaks first
                for i in range(search_end - 1, search_start - 1, -1):
                    if i + 1 < text_length and text[i:i+2] == '\n\n':
                        break_candidates.append(i + 2)
                        break
                
                # Look for sentence breaks
                if not break_candidates:
                    for i in range(search_end - 1, search_start - 1, -1):
                        if text[i] in '.!?' and i + 1 < text_length and text[i + 1] in ' \n':
                            break_candidates.append(i + 1)
                            break
                
                # Look for line breaks
                if not break_candidates:
                    for i in range(search_end - 1, search_start - 1, -1):
                        if text[i] == '\n':
                            break_candidates.append(i + 1)
                            break
                
                # Use the best break point found, or default to fixed position
                if break_candidates:
                    break_point = break_candidates[0]
                
                # Extract chunk
                chunk_text = text[start:break_point].strip()
                
                # Add header context for markdown files
                if headers_context and file_type == '.md':
                    chunk_text = headers_context + "\n\n" + chunk_text
                
                if chunk_text:
                    chunks.append(chunk_text)
                
                # Move start position (with overlap)
                start = break_point - chunk_overlap if break_point > chunk_overlap else break_point
            
            logger.info(f"📊 Fixed-size split: {len(chunks)} chunks created")
            for i, chunk in enumerate(chunks):
                logger.debug(f"   Chunk {i+1}: {len(chunk)} chars")
            
            return chunks
            
        except Exception as e:
            logger.error(f"❌ Fixed-size splitting failed: {e}")
            return self._simple_split(text)
    
    def _extract_document_headers(self, text: str) -> str:
        """Extract main document headers for context"""
        lines = text.split('\n')
        headers = []
        
        for line in lines[:20]:  # Check first 20 lines for main headers
            line = line.strip()
            if line.startswith('#') and len(line.split()) <= 10:  # Reasonable header length
                # Only keep H1 and H2 headers for context
                if line.startswith('##') and not line.startswith('###'):
                    headers.append(line)
                elif line.startswith('#') and not line.startswith('##'):
                    headers.append(line)
            
            # Stop if we find content after headers
            if not line.startswith('#') and line and len(headers) > 0:
                break
        
        return '\n'.join(headers) if headers else ""
    
    def _standard_split(self, text: str) -> List[str]:
        """Standard LangChain splitting without header context"""
        try:
            chunks = self.text_splitter.split_text(text)
            chunks = [chunk.strip() for chunk in chunks if chunk.strip()]
            
            logger.info(f"🎯 [CHUNKER] Standard splitting generated {len(chunks)} chunks")
            
            # Log chunk size statistics
            if chunks:
                chunk_sizes = [len(chunk) for chunk in chunks]
                avg_size = sum(chunk_sizes) / len(chunk_sizes)
                min_size = min(chunk_sizes)
                max_size = max(chunk_sizes)
                
                logger.debug(f"📊 [CHUNKER] Chunk stats: avg={avg_size:.0f}, min={min_size}, max={max_size}")
            
            return chunks
            
        except Exception as e:
            logger.error(f"💥 [CHUNKER] Standard splitting failed: {e}")
            return self._simple_split(text)
    
    def _has_headers(self, text: str) -> bool:
        """Check if text has header-like structures"""
        # Check for markdown headers
        if re.search(r'^#{1,6}\s+.+$', text, re.MULTILINE):
            return True
        
        # Check for numbered headers (1. 2. etc.)
        if re.search(r'^\d+\.\s+.+$', text, re.MULTILINE):
            return True
        
        # Check for ALL CAPS headers
        if re.search(r'^[A-Z][A-Z\s]{3,}$', text, re.MULTILINE):
            return True
        
        return False
    
    def _extract_headers(self, text: str) -> List[dict]:
        """Extract headers and their positions from text"""
        headers = []
        lines = text.split('\n')
        position = 0
        
        for line in lines:
            line_stripped = line.strip()
            
            # Markdown headers
            md_match = re.match(r'^(#{1,6})\s+(.+)$', line_stripped)
            if md_match:
                level = len(md_match.group(1))
                title = md_match.group(2)
                headers.append({
                    'level': level,
                    'title': title,
                    'position': position,
                    'type': 'markdown'
                })
            
            # Numbered headers
            elif re.match(r'^\d+\.\s+.+$', line_stripped):
                headers.append({
                    'level': 2,  # Treat as level 2
                    'title': line_stripped,
                    'position': position,
                    'type': 'numbered'
                })
            
            # ALL CAPS headers
            elif re.match(r'^[A-Z][A-Z\s]{3,}$', line_stripped) and len(line_stripped) < 100:
                headers.append({
                    'level': 1,  # Treat as level 1
                    'title': line_stripped,
                    'position': position,
                    'type': 'caps'
                })
            
            position += len(line) + 1  # +1 for newline
        
        return headers
    
    def _find_relevant_headers(self, headers: List[dict], chunk_position: int) -> List[dict]:
        """Find headers that should provide context for a chunk at given position"""
        relevant_headers = []
        
        # Find headers before this chunk position
        for header in reversed(headers):  # Start from most recent
            if header['position'] <= chunk_position:
                # Add this header and any higher-level headers above it
                relevant_headers.insert(0, header)
                
                # Look for higher-level headers
                current_level = header['level']
                for prev_header in reversed(headers):
                    if (prev_header['position'] < header['position'] and 
                        prev_header['level'] < current_level):
                        relevant_headers.insert(0, prev_header)
                        current_level = prev_header['level']
                        if current_level == 1:  # Found top level
                            break
                break
        
        # Remove duplicates while preserving order
        seen = set()
        unique_headers = []
        for header in relevant_headers:
            key = (header['position'], header['title'])
            if key not in seen:
                seen.add(key)
                unique_headers.append(header)
        
        return unique_headers
    
    def _build_header_context(self, metadata: dict) -> str:
        """Build header context string from markdown metadata"""
        context_parts = []
        
        # Sort by header level
        for i in range(1, 7):
            header_key = f"Header {i}"
            if header_key in metadata:
                level_prefix = "#" * i
                context_parts.append(f"{level_prefix} {metadata[header_key]}")
        
        return "\n".join(context_parts) + "\n\n" if context_parts else ""
    
    def _build_text_header_context(self, headers: List[dict]) -> str:
        """Build header context string from extracted headers"""
        if not headers:
            return ""
        
        context_parts = []
        for header in headers:
            if header['type'] == 'markdown':
                level_prefix = "#" * header['level']
                context_parts.append(f"{level_prefix} {header['title']}")
            else:
                context_parts.append(header['title'])
        
        return "\n".join(context_parts) if context_parts else ""
    
    def _simple_split(self, text: str) -> List[str]:
        """
        Simple fallback text splitting method
        
        Args:
            text: Text to split
            
        Returns:
            List of text chunks
        """
        logger.debug(f"🔪 [CHUNKER] Simple fallback split for {len(text)} chars")
        
        if len(text) <= self.chunk_size:
            return [text.strip()] if text.strip() else []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to find a good break point (sentence ending)
            if end < len(text):
                # Look for sentence endings within the last 100 characters
                search_start = max(start + self.chunk_size - 100, start)
                search_text = text[search_start:end]
                
                # Find sentence endings
                sentence_endings = [m.end() for m in re.finditer(r'[.!?]\s+', search_text)]
                if sentence_endings:
                    # Use the last sentence ending
                    end = search_start + sentence_endings[-1]
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
                logger.debug(f"📦 [CHUNKER] Simple chunk: {len(chunk)} chars")
            
            # Move start position with overlap
            start = max(end - self.chunk_overlap, start + 1)
            
            # Prevent infinite loop
            if start >= len(text):
                break
        
        logger.debug(f"✅ [CHUNKER] Simple split created {len(chunks)} chunks")
        return chunks
    
    def process_and_save_to_milvus(self, file_path: str, filename: str, milvus_service) -> bool:
        """
        Process file and save extracted content to Milvus
        
        Args:
            file_path: Path to the uploaded file
            filename: Original filename
            milvus_service: Milvus service instance
            
        Returns:
            True if successful, False otherwise
        """
        logger.info(f"🚀 [MILVUS-PROCESSOR] Starting Milvus processing for: {filename}")
        
        try:
            # Extract content from file
            logger.info(f"📄 [MILVUS-PROCESSOR] Step 1: Extracting content from {filename}")
            doc = self.process_uploaded_file(file_path, filename)
            if not doc:
                logger.error(f"❌ [MILVUS-PROCESSOR] Failed to extract content from {filename}")
                return False
            
            logger.info(f"✅ [MILVUS-PROCESSOR] Content extracted successfully")
            
            # Split content into chunks with context preservation
            logger.info(f"✂️ [MILVUS-PROCESSOR] Step 2: Splitting content into chunks with header preservation")
            chunks = self._split_text(doc["content"], doc["file_type"])
            logger.info(f"📊 [MILVUS-PROCESSOR] Generated {len(chunks)} chunks from content")
            
            if not chunks:
                logger.warning(f"⚠️ [MILVUS-PROCESSOR] No chunks generated from {filename}")
                return False
            
            # Prepare documents for Milvus
            logger.info(f"🔧 [MILVUS-PROCESSOR] Step 3: Preparing documents for Milvus")
            milvus_docs = []
            for chunk_idx, chunk in enumerate(chunks):
                milvus_doc = {
                    "file_name": doc["file_name"],
                    "chunk_id": chunk_idx,
                    "content": chunk,
                    "title": doc["title"],
                    "section": f"Chunk {chunk_idx + 1}"
                }
                milvus_docs.append(milvus_doc)
                logger.debug(f"📝 [MILVUS-PROCESSOR] Chunk {chunk_idx}: {len(chunk)} characters")
            
            logger.info(f"✅ [MILVUS-PROCESSOR] Prepared {len(milvus_docs)} documents for Milvus")
            
            # Save to Milvus
            logger.info(f"💾 [MILVUS-PROCESSOR] Step 4: Saving to Milvus database")
            logger.info(f"🔌 [MILVUS-PROCESSOR] Milvus service available: {milvus_service is not None}")
            
            success = milvus_service.insert_documents(milvus_docs)
            if success:
                logger.info(f"🎉 [MILVUS-PROCESSOR] Successfully saved {len(milvus_docs)} chunks from {filename} to Milvus")
                
                # Verify data was saved
                stats = milvus_service.get_collection_stats()
                logger.info(f"📈 [MILVUS-PROCESSOR] Collection now has {stats} total entities")
                return True
            else:
                logger.error(f"❌ [MILVUS-PROCESSOR] Failed to save chunks from {filename} to Milvus")
                return False
                
        except Exception as e:
            logger.error(f"💥 [MILVUS-PROCESSOR] Error processing and saving {filename} to Milvus: {e}", exc_info=True)
            return False
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize text
        
        Args:
            text: Raw text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove markdown formatting
        text = re.sub(r'[#*`_]+', '', text)
        
        # Trim
        text = text.strip()
        
        return text
